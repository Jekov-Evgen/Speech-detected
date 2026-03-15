import os
import sys
import shutil
import asyncio
from io import BytesIO
from datetime import datetime
import uvicorn
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Cm

from database import init_db, add_record, update_transcript, get_all_records, delete_record
from transcriber import transcribe, progress_store, model_status


def _base_dir():
    """Директория с ресурсами (шрифты). В PyInstaller — sys._MEIPASS."""
    if getattr(sys, 'frozen', False):
        return sys._MEIPASS
    return os.path.dirname(__file__)


def _data_dir():
    """Записываемая директория для данных (uploads, БД, .env)."""
    d = os.environ.get("SPEECHDETECT_DATA_DIR")
    if d:
        os.makedirs(d, exist_ok=True)
        return d
    return os.path.dirname(__file__)


FONTS_DIR = os.path.join(_base_dir(), "fonts")
UPLOAD_DIR = os.path.join(_data_dir(), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

app = FastAPI(title="SpeechDetect API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
def startup():
    init_db()


@app.get("/api/health")
def health():
    from transcriber import model_status
    return {"status": "ok", "model_status": model_status}


@app.post("/api/upload")
async def upload_audio(file: UploadFile = File(...)):
    allowed = {".wav", ".mp3", ".ogg", ".flac", ".m4a"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed:
        raise HTTPException(400, f"Формат {ext} не поддерживается")

    filepath = os.path.join(UPLOAD_DIR, file.filename)

    base, extension = os.path.splitext(filepath)
    counter = 1
    while os.path.exists(filepath):
        filepath = f"{base}_{counter}{extension}"
        counter += 1

    with open(filepath, "wb") as f:
        shutil.copyfileobj(file.file, f)

    filesize = os.path.getsize(filepath)
    record_id = add_record(file.filename, filepath, filesize)

    return {"id": record_id, "filename": file.filename, "filepath": filepath, "filesize": filesize}


@app.post("/api/transcribe/{record_id}")
async def transcribe_audio(record_id: int):
    records = get_all_records()
    record = next((r for r in records if r["id"] == record_id), None)
    if not record:
        raise HTTPException(404, "Запись не найдена")

    if not os.path.exists(record["filepath"]):
        raise HTTPException(404, "Аудиофайл не найден на диске")

    try:
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(None, transcribe, record["filepath"], record_id)
    except Exception as e:
        progress_store.pop(record_id, None)
        raise HTTPException(500, f"Ошибка распознавания: {e}")

    update_transcript(record_id, text)
    progress_store.pop(record_id, None)
    return {"id": record_id, "transcript": text}


@app.get("/api/progress/{record_id}")
def get_progress(record_id: int):
    info = progress_store.get(record_id)
    if info is None:
        return {"progress": 0.0, "status": "unknown"}
    return info


@app.get("/api/export/pdf/{record_id}")
def export_pdf(record_id: int):
    records = get_all_records()
    record = next((r for r in records if r["id"] == record_id), None)
    if not record:
        raise HTTPException(404, "Запись не найдена")
    if not record.get("transcript"):
        raise HTTPException(400, "Нет текста для экспорта")

    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("DejaVu", "", os.path.join(FONTS_DIR, "DejaVuSans.ttf"))
    pdf.add_font("DejaVu", "B", os.path.join(FONTS_DIR, "DejaVuSans-Bold.ttf"))

    pdf.set_font("DejaVu", "B", 16)
    pdf.cell(0, 12, "Результат распознавания речи", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("DejaVu", "", 10)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 8, f"Файл: {record['filename']}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 8, f"Дата: {record.get('created_at', '')}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)

    pdf.set_text_color(0, 0, 0)
    pdf.set_font("DejaVu", "", 12)
    pdf.multi_cell(0, 7, record["transcript"])

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)

    safe_name = os.path.splitext(record["filename"])[0]
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.pdf"'},
    )


@app.get("/api/export/docx/{record_id}")
def export_docx(record_id: int):
    records = get_all_records()
    record = next((r for r in records if r["id"] == record_id), None)
    if not record:
        raise HTTPException(404, "Запись не найдена")
    if not record.get("transcript"):
        raise HTTPException(400, "Нет текста для экспорта")

    doc = Document()

    title = doc.add_heading("Результат распознавания речи", level=1)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(2)
    run = meta.add_run(f"Файл: {record['filename']}")
    run.font.size = Pt(10)
    run.font.color.rgb = None

    meta2 = doc.add_paragraph()
    meta2.paragraph_format.space_after = Pt(12)
    run2 = meta2.add_run(f"Дата: {record.get('created_at', '')}")
    run2.font.size = Pt(10)

    body = doc.add_paragraph(record["transcript"])
    body.style.font.size = Pt(12)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = os.path.splitext(record["filename"])[0]
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
    )


@app.get("/api/history")
def history():
    return get_all_records()


@app.delete("/api/history/{record_id}")
def delete_history(record_id: int):
    records = get_all_records()
    record = next((r for r in records if r["id"] == record_id), None)
    if record and os.path.exists(record["filepath"]):
        os.remove(record["filepath"])
    delete_record(record_id)
    return {"ok": True}


if __name__ == "__main__":
    uvicorn.run(app, host="127.0.0.1", port=8765)